# def login_view(request):
#     if request.method == "POST":
#         username = request.POST.get('username')
#         password = request.POST.get('password')

#         if not username or not password:
#             messages.error(request, 'Por favor, preencha ambos os campos.')
#             return render(request, 'login.html')

#         user = authenticate(request, username=username, password=password)

#         if user is not None:
#             login(request, user)
#             return redirect('home')
#         else:
#             messages.error(request, 'Nome de usuário ou senha incorretos.')

#     return render(request, 'login.html')


# ##################################################################

from .models import PassagemDeServico, CautelaDeArmamento, RegistroDescautelamento
import pypandoc
from docx import Document
from .models import CautelaDeArmamento, PassagemDeServico, User
from django.shortcuts import render
import os
from django.utils.html import format_html
import base64
from docx.oxml.ns import nsmap
from django.conf import settings
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import asyncio
from docx.shared import Inches
# from pyppeteer import launch
from django.template.loader import render_to_string
from django.http import HttpResponse
from datetime import datetime
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponseBadRequest, HttpResponse
from django.contrib.auth.models import User
from django.template.loader import render_to_string
from django.db.models import Sum
from django.db import transaction
from django.core.paginator import Paginator
from django.urls import reverse
from django.core.mail import send_mail
from django.contrib import messages
from django.db import IntegrityError
from django.utils.html import strip_tags
from html import escape
from django.utils.html import format_html
import json
from docx.oxml import parse_xml
from .models import PassagemDeServico
from weasyprint import HTML
from weasyprint.css import CSS
from docx2pdf import convert
import logging
from collections import defaultdict
import subprocess
from django.utils import timezone
from .models import (
    Categoria,
    CautelaDeArmamento,
    Subcategoria,
    Policial,
    DescautelasCa,
    CategoriaMunicao,
    SubcategoriaMunicao,
    CautelaDeMunicoes,
    RegistroCautelaCompleta,
    RegistroDescautelamento,
    PassagemDeServico
)


def index(request):
    return render(request, 'cautelaarmamento/index.html')


def cautela_de_armamento_view(request):

    ############## PROCESSAMENTO DE REQUISIÇÃO POST ##########

    if request.method == 'POST':
        # Obtém os dados do formulário
        policial_id = request.POST.get('policial')
        tipo_servico = request.POST.get('tipo_servico')
        categorias_armamento = request.POST.getlist('categorias[]')
        subcategorias_armamento = request.POST.getlist('subcategorias[]')
        categorias_municao = request.POST.getlist('categoria_municoes[]')
        subcategorias_municao = request.POST.getlist('subcategoria_municoes[]')
        quantidades_municao = request.POST.getlist('quantidade[]')


############## VALIDAÇÃO DOS CAMPOS OBRIGATÓRIOS ##########

        # Verifica se os campos obrigatórios foram preenchidos
        if not policial_id or not tipo_servico:
            return JsonResponse({'error': 'Os campos Nome do Policial e Tipo de Serviço são obrigatórios.'}, status=400)

        policial = get_object_or_404(Policial, id=policial_id, autorizacao_arma=True)
        armeiro = request.user

############## PROCESSAMENTO DO ARMAMENTO ##########

        try:
            with transaction.atomic():
                # Inicializa variáveis para munição e armamento
                categoria_armamento_nome = None
                subcategoria_armamento_nome = None
                categoria_municao_nome = None
                subcategoria_municao_nome = None
                quantidade_total_municao = 0

                # Processa o armamento, se presente

                if categorias_armamento and subcategorias_armamento:
                    for i in range(len(categorias_armamento)):
                        categoria_id = categorias_armamento[i]
                        subcategoria_id = subcategorias_armamento[i]
                        if categoria_id and subcategoria_id:
                            categoria_armamento = get_object_or_404(
                                Categoria, id=categoria_id)
                            subcategoria_armamento = get_object_or_404(
                                Subcategoria, id=subcategoria_id)

                            # Salva registro no modelo de Cautela de Armamento
                            CautelaDeArmamento.objects.create(
                                policial=policial,
                                tipo_servico=tipo_servico,
                                categoria=categoria_armamento,
                                subcategoria=subcategoria_armamento,
                                armeiro=armeiro
                            )

                            subcategoria_armamento.situacao = 'cautelada'
                            subcategoria_armamento.save()

                            # Registro na nova model com todos os detalhes para cada linha
                            RegistroCautelaCompleta.objects.create(
                                policial=policial,
                                tipo_servico=tipo_servico,
                                categoria_armamento=categoria_armamento.nome,  # Aqui pode permanecer igual
                                subcategoria_armamento=subcategoria_armamento.descricao_completa,  # Mudança feita aqui
                                categoria_municao=None,
                                subcategoria_municao=None,
                                quantidade_municao=0,
                                armeiro=armeiro
                            )
                # Processa as munições

############## PROCESSAMENTO DA MUNIÇÃO ##########

                if categorias_municao and subcategorias_municao and quantidades_municao:
                    for i in range(len(categorias_municao)):
                        categoria_id = categorias_municao[i]
                        subcategoria_id = subcategorias_municao[i]
                        quantidade = quantidades_municao[i]
                        if categoria_id and subcategoria_id and quantidade:
                            categoria_municao = get_object_or_404(
                                CategoriaMunicao, id=categoria_id)
                            subcategoria_municao = get_object_or_404(
                                SubcategoriaMunicao, id=subcategoria_id)

                            quantidade = int(quantidade)
                            if quantidade > subcategoria_municao.total_de_municoes:
                                return JsonResponse({'error': 'A quantidade solicitada excede o total disponível.'}, status=400)

                            # Debita o valor no banco de dados
                            subcategoria_municao.total_de_municoes -= quantidade
                            subcategoria_municao.save()

                            # Salva registro no modelo de Cautela de Munição
                            CautelaDeMunicoes.objects.create(
                                policial=policial,
                                categoria=categoria_municao,
                                subcategoria=subcategoria_municao,
                                quantidade=quantidade,
                                armeiro=armeiro,
                            )

                            # Registro na nova model com todos os detalhes para cada linha - Munição
                            RegistroCautelaCompleta.objects.create(
                                policial=policial,
                                tipo_servico=tipo_servico,
                                categoria_armamento=None,
                                subcategoria_armamento=None,
                                categoria_municao=categoria_municao.nome,  # Aqui pode permanecer igual
                                # Atualize aqui para o atributo correto
                                subcategoria_municao=subcategoria_municao.nome,
                                quantidade_municao=quantidade,
                                armeiro=armeiro
                            )

############## GERAÇÃO DE RELATÓRIO E ENVIO DE EMAIL ##########

        # Após o processamento bem-sucedido, gere o relatório e envie o email
            try:
                # Obtenha a data e hora atual
                data_hora_atual = timezone.now()
                with transaction.atomic():
                    # Seu código de processamento de armamento e munição...

                    # Aqui você gera o relatório para o email
                    email_subject = 'Relatório de Cautela de Armamento e Munição'
                    email_recipients = ['mendesjunior2007@hotmail.com']

                    # Crie o conteúdo do email com base nos dados processados
                    # Crie o conteúdo do email com base nos dados processados
                    context = {
                        'policial': policial.nome_completo,
                        'tipo_servico': tipo_servico,
                        'armamentos': [
                            {
                                'categoria': get_object_or_404(Categoria, id=categorias_armamento[i]).nome,
                                'subcategoria': get_object_or_404(Subcategoria, id=subcategorias_armamento[i]).descricao_completa
                            }
                            for i in range(len(categorias_armamento))
                        ],
                        'municoes': [
                            {
                                'categoria': get_object_or_404(CategoriaMunicao, id=categorias_municao[i]).nome,
                                'subcategoria': get_object_or_404(SubcategoriaMunicao, id=subcategorias_municao[i]).nome,
                                'quantidade': quantidades_municao[i]
                            }
                            # Verifica se existem munições
                            for i in range(len(categorias_municao)) if categorias_municao[i] and subcategorias_municao[i]
                        ],
                        # Formata a data e hora
                        'data_hora': data_hora_atual.strftime('%d/%m/%Y %H:%M'),
                        # Obtém o nome completo do usuário logado
                        'usuario': request.user.get_full_name()
                    }

                    # Use um template HTML para o email
                    html_message = render_to_string(
                        'email/relatorio_cautela.html', context)
                    plain_message = strip_tags(html_message)

                    # Envie o email
                    send_mail(
                        subject=email_subject,
                        message=plain_message,
                        from_email='seuemail@gmail.com',
                        recipient_list=email_recipients,
                        html_message=html_message,
                        fail_silently=False,
                    )

############## EXCEÇÃO GERAL E REDIRECIONAMENTO ##########

            except Exception as e:
                return JsonResponse({'error': str(e)}, status=500)

            return redirect('sucesso')
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

        return redirect('sucesso')

############## EXIBIÇÃO DO FORMULÁRIO E CONTEXTO ##########

    # Renderiza o template de cautela
    policiais = Policial.objects.filter(autorizacao_arma=True)
    tipos_servico = CautelaDeArmamento.SERVICO_CHOICES
    categorias_servico = Categoria.objects.all()
    subcategoria_armamento = Subcategoria.objects.all()
    categoria_municoes = CategoriaMunicao.objects.all()
    subcategoria_municoes = SubcategoriaMunicao.objects.all()

    # Calcula a quantidade total de munições
    quantidade_total = CautelaDeMunicoes.objects.aggregate(
        total=Sum('quantidade'))['total']
    if quantidade_total is None:
        quantidade_total = 0

    context = {
        'policiais': policiais,
        'categorias': categorias_servico,
        'tipos_servico': tipos_servico,
        'subcategorias_armamento': subcategoria_armamento,
        'categorias_municoes': categoria_municoes,
        'subcategorias_municoes': subcategoria_municoes,
        'quantidade_total': quantidade_total
    }

    return render(request, 'armamento/cautela.html', context)


############################# AJAX ###################################
######################################################################
# busca todas as subcategorias relacionadas a armamento
def get_subcategorias_armamento(request, categoria_id):
    try:
        categoria = Categoria.objects.get(id=categoria_id)
        subcategorias = categoria.subcategorias_armamento.filter(
            situacao='disponivel')
    except Categoria.DoesNotExist:
        return JsonResponse({'error': 'Categoria não encontrada'}, status=404)

    # Altere o valor aqui de subcategoria.nome para subcategoria.descricao_completa
    data = [{"id": subcategoria.id, "nome": subcategoria.descricao_completa}
            for subcategoria in subcategorias]
    return JsonResponse(data, safe=False)

# busca todas as subcategorias relacionadas a munição


def get_subcategorias_municao(request, categoria_id):
    categoria = get_object_or_404(CategoriaMunicao, id=categoria_id)
    subcategorias = categoria.subcategorias.all()
    # Altere "nome" para "descricao_completa" abaixo:
    data = [{"id": subcategoria.id, "nome": subcategoria.nome}
            for subcategoria in subcategorias]
    return JsonResponse(data, safe=False)

# busca todas as quantidades de munições


def obter_quantidade_total(request, subcategoria_id):
    try:
        subcategoria = get_object_or_404(
            SubcategoriaMunicao, id=subcategoria_id)
        total_municoes = subcategoria.total_de_municoes
        return JsonResponse({'quantidade': total_municoes})
    except SubcategoriaMunicao.DoesNotExist:
        return JsonResponse({'error': 'Subcategoria não encontrada'}, status=404)


def listar_registros_cautela(request):
    # Recupera todos os registros de cautela do banco de dados
    registros = RegistroCautelaCompleta.objects.all()

    # Passa os registros para o template
    context = {
        'registros': registros
    }

    return render(request, 'armamento/descautela.html', context)

######################################################################
######################################################################

# def inventario_equipamentos(request):
#     equipamentos = CautelaDeArmamento.objects.all()
#     context = {'equipamentos': equipamentos}
#     return render(request, 'inventario.html', context)

# def descautelar_armamento(request):
#     if request.method == 'POST':
#         cautela_id = request.POST.get('cautela_id')
#         cautela = get_object_or_404(CautelaDeArmamento, id=cautela_id)
#         cautela.delete()
#         return JsonResponse({'success': True})
#     return JsonResponse({'success': False})

# def listar_armamentos(request):
#     armamentos = CautelaDeArmamento.objects.all()
#     context = {'armamentos': armamentos}
#     return render(request, 'listar_armamentos.html', context)


def cadastrar_pessoa(request):
    if request.method == 'POST':
        nome_completo = request.POST.get('nome_completo')
        nome_guerra = request.POST.get('nome_guerra')
        posto_graduacao = request.POST.get('posto_graduacao')
        matricula = request.POST.get('matricula')
        rgpm = request.POST.get('rgpm')
        lotacao = request.POST.get('lotacao')
        data_nascimento = request.POST.get('data_nascimento')
        cpf = request.POST.get('cpf')

        try:
            # Verificar se já existe um policial com o mesmo CPF
            if Policial.objects.filter(cpf=cpf).exists():
                return render(request, 'armamento/cadastrar_pessoa.html', {
                    'error': 'Um policial com este CPF já está cadastrado.'
                })

            # Criar e salvar o novo policial
            policial = Policial(
                nome_completo=nome_completo,
                nome_guerra=nome_guerra,
                posto_graduacao=posto_graduacao,
                matricula=matricula,
                rgpm=rgpm,
                lotacao=lotacao,
                data_nascimento=data_nascimento,
                cpf=cpf
            )
            policial.save()

            # Redirecionar para a página de sucesso
            return redirect('sucesso')

        except IntegrityError:
            # Renderizar a página com uma mensagem de erro caso ocorra um problema de integridade
            return render(request, 'armamento/cadastrar_pessoa.html', {
                'error': 'Ocorreu um erro ao tentar salvar o policial. Verifique os dados e tente novamente.'
            })

    # Renderizar a página de cadastro se a requisição for GET
    return render(request, 'armamento/cadastrar_pessoa.html')


def sucesso_view(request):
    # Função de visualização de sucesso após o cadastro de uma pessoa
    return render(request, 'armamento\sucesso.html')

#############################################################
###########################################################


def descautelar_sa(request):
    if request.method == 'POST':
        registro_id = request.POST.get('registro_id')

        # Buscando o registro no banco de dados com base no ID
        registro = get_object_or_404(RegistroCautelaCompleta, id=registro_id)

        # Obter a data e hora atuais para o descautelamento
        data_hora_atual = timezone.now()

        # Verificar se a categoria de armamento está presente
        if registro.categoria_armamento:
            # Buscar a subcategoria de armamento associada ao registro
            subcategorias = Subcategoria.objects.filter(
                descricao_completa=registro.subcategoria_armamento)

            if subcategorias.exists():
                # Se houver múltiplas subcategorias, pegar a primeira
                subcategoria = subcategorias.first()

                # Alterar o campo situacao para 'disponivel'
                subcategoria.situacao = 'disponivel'
                subcategoria.save()

                # Registrar o descautelamento no novo modelo
                RegistroDescautelamento.objects.create(
                    data_hora_cautela=data_hora_atual,
                    policial=registro.policial,  # Atribuindo a instância do policial
                    tipo_servico=registro.tipo_servico,
                    categoria_armamento=registro.categoria_armamento,
                    subcategoria_armamento=registro.subcategoria_armamento,
                    situacao_armamento='disponível',  # Situação após descautela
                    armeiro=request.user,  # Usuário que realizou o descautelamento
                    observacao='Descautela de armamento realizada automaticamente.'
                )

                # Imprimir o valor da categoria armamento no terminal
                print(
                    f"Categoria de Armamento: {registro.categoria_armamento}")
                print(
                    f"Subcategoria '{subcategoria.descricao_completa}' alterada para situação: {subcategoria.situacao}")
            else:
                # Se não houver subcategorias, lidar com a situação (opcional)
                print("Nenhuma subcategoria encontrada.")

        # Caso a categoria de armamento seja None, trabalhar com munição
        elif registro.categoria_armamento is None and registro.quantidade_municao > 0:
            # Buscar a subcategoria de munição associada ao registro
            subcategoria_municao = get_object_or_404(
                SubcategoriaMunicao, nome=registro.subcategoria_municao)

            # Atualizar o total de munições
            subcategoria_municao.total_de_municoes += registro.quantidade_municao
            subcategoria_municao.save()

            # Registrar o descautelamento no novo modelo
            RegistroDescautelamento.objects.create(
                data_hora_cautela=data_hora_atual,
                policial=registro.policial,  # Atribuindo a instância do policial
                tipo_servico=registro.tipo_servico,
                categoria_municao=registro.categoria_municao,
                subcategoria_municao=registro.subcategoria_municao,
                quantidade_municao=registro.quantidade_municao,
                situacao_armamento='N/A',  # Não se aplica a munições
                armeiro=request.user,  # Usuário que realizou o descautelamento
                observacao='Descautela de munição realizada automaticamente.'
            )

            # Imprimir o valor da quantidade de munições no terminal
            print(f"Quantidade de Munição: {registro.quantidade_municao}")
            print(
                f"Subcategoria de Munição '{subcategoria_municao.nome}' agora tem {subcategoria_municao.total_de_municoes} munições.")

        # Após o processo, excluir o registro de cautela completa
        registro.delete()
        print(
            f"Registro de cautela completa {registro_id} excluído do banco de dados.")

        # Retornar uma resposta de sucesso
        return JsonResponse({'status': 'success', 'registro_id': registro_id})

    # Caso não seja POST, retornar uma resposta de erro
    return JsonResponse({'status': 'failed', 'message': 'Invalid request method'})


def descautelar_ca(request):
    if request.method == 'POST':
        try:
            # Obter os dados diretamente de request.POST
            registro_id = request.POST.get('registro_id')
            situacao_armamento = request.POST.get('situacao')
            observacao = request.POST.get('observacao', '')
            observacoes_input = request.POST.get('observacoes', '')

            # Obter a quantidade de munição digitada
            quantidade_digitada = int(
                request.POST.get('quantidade_municao', 0))

            # Exibir a quantidade de munição digitada no terminal
            print(
                f"Quantidade de Munição Digitada pelo Usuário: {quantidade_digitada}")

            # Obtém o registro específico usando o ID fornecido
            registro = get_object_or_404(
                RegistroCautelaCompleta, pk=registro_id)

            # Caso a subcategoria de munição seja None, continuar com o fluxo normal
            if registro.subcategoria_municao is None:
                policial = registro.policial
                armeiro_descautela = request.user
                data_hora_atual = timezone.now()

                categoria_municao = registro.categoria_municao if registro.categoria_municao else "Categoria Padrão"
                subcategoria_municao = registro.subcategoria_municao if registro.subcategoria_municao else "Subcategoria Padrão"

                print(
                    f"Quantidade Original de Munição no Registro: {registro.quantidade_municao}")

                # Criação do novo registro de descautelamento
                novo_descautelamento = DescautelasCa.objects.create(
                    data_hora_cautela=registro.data_hora if hasattr(
                        registro, 'data_hora') else timezone.now(),
                    policial=registro.policial,
                    tipo_servico=registro.tipo_servico,
                    categoria_armamento=registro.categoria_armamento,
                    subcategoria_armamento=registro.subcategoria_armamento,
                    categoria_municao=categoria_municao,
                    subcategoria_municao=subcategoria_municao,
                    quantidade_municao=quantidade_digitada,
                    situacao_armamento=situacao_armamento,
                    observacao=observacao,
                    armeiro=registro.armeiro,
                    armeiro_descautela=request.user,
                    data_descautelamento=timezone.now().date(),
                    hora_descautelamento=timezone.now().time()
                )

                print("Dados do Novo Registro de Descautelamento:")
                print([
                    novo_descautelamento.data_hora_cautela,
                    novo_descautelamento.policial,
                    novo_descautelamento.tipo_servico,
                    novo_descautelamento.categoria_armamento,
                    novo_descautelamento.subcategoria_armamento,
                    novo_descautelamento.categoria_municao,
                    novo_descautelamento.subcategoria_municao,
                    novo_descautelamento.quantidade_municao,
                    novo_descautelamento.situacao_armamento,
                    novo_descautelamento.observacao,
                    novo_descautelamento.armeiro,
                    novo_descautelamento.armeiro_descautela,
                    novo_descautelamento.data_descautelamento,
                    novo_descautelamento.hora_descautelamento
                ])

                if registro.subcategoria_armamento:
                    subcategoria = get_object_or_404(
                        Subcategoria, descricao_completa=registro.subcategoria_armamento)

                    subcategoria.situacao = situacao_armamento
                    subcategoria.save()

                registro.delete()
                return JsonResponse({'success': True, 'message': 'Descautelamento realizado com sucesso.', 'registro_id': registro_id})

            else:
                quantidade_municao_disponivel = registro.quantidade_municao
                print(
                    f"Quantidade de Munição Disponível: {quantidade_municao_disponivel}")

                categoria_municao = registro.categoria_municao if registro.categoria_municao else "Categoria Padrão"
                subcategoria_municao = registro.subcategoria_municao if registro.subcategoria_municao else "Subcategoria Padrão"

                print(f"Categoria de Munição: {categoria_municao}")
                print(f"Subcategoria de Munição: {subcategoria_municao}")
                novo_descautelamento = DescautelasCa.objects.create(
                    data_hora_cautela=registro.data_hora if hasattr(
                        registro, 'data_hora') else timezone.now(),
                    policial=registro.policial,
                    tipo_servico=registro.tipo_servico,
                    categoria_armamento=registro.categoria_armamento,
                    subcategoria_armamento=registro.subcategoria_armamento,
                    categoria_municao=categoria_municao,
                    subcategoria_municao=subcategoria_municao,
                    quantidade_municao=quantidade_digitada,
                    situacao_armamento=situacao_armamento,
                    observacao=observacao,
                    observacoes=observacoes_input,  # Adicione esta linha
                    armeiro=registro.armeiro,
                    armeiro_descautela=request.user,
                    data_descautelamento=timezone.now().date(),
                    hora_descautelamento=timezone.now().time()
                )

                print("Novo registro de descautelamento criado com sucesso.")
                print(f"ID do novo descautelamento: {novo_descautelamento.id}")

                quantidade_restante = quantidade_municao_disponivel - quantidade_digitada
                print(
                    f"Quantidade Restante após Descautela: {quantidade_restante}")

                if quantidade_restante >= 0:
                    subcategoria_municao_obj = get_object_or_404(
                        SubcategoriaMunicao, nome=registro.subcategoria_municao)
                    print(
                        f"Total de Munições Antes da Atualização: {subcategoria_municao_obj.total_de_municoes}")

                    if subcategoria_municao_obj.total_de_municoes is None:
                        subcategoria_municao_obj.total_de_municoes = 0

                    subcategoria_municao_obj.total_de_municoes += quantidade_restante
                    subcategoria_municao_obj.save()

                    print(
                        f"Quantidade Atualizada de Munições na Subcategoria: {subcategoria_municao_obj.total_de_municoes}")

                registro.delete()
                return JsonResponse({'success': True, 'limite_municao': quantidade_municao_disponivel, 'quantidade_restante': quantidade_restante})

        except Exception as e:
            print(
                f"Erro durante a criação do registro de descautelamento: {str(e)}")
            return JsonResponse({'success': False, 'error': str(e)})

    # Caso o método não seja POST, renderize uma página de erro ou mensagem adequada
    return JsonResponse({'success': False, 'message': 'Método não permitido.'}, status=405)


def descautelar_municao_ca(request):
    if request.method == 'POST':
        # Verifique se isto aparece no terminal
        print("Função descautelar_municao_ca foi chamada")
        # Seus dados aqui...
        registro_id = request.POST.get('registro_id')
        quantidade_municao = request.POST.get('quantidade_municao')
        situacao = request.POST.get('situacao')
        quantidade_atual = request.POST.get('quantidade_atual')

        # Exibindo os dados no terminal
        print(f"Registro ID: {registro_id}")
        print(f"Quantidade de Munição: {quantidade_municao}")
        print(f"Situação do Armamento: {situacao}")
        print(f"Quantidade Atual: {quantidade_atual}")

        # Retornar uma resposta de sucesso como JSON
        return JsonResponse({'success': True})

    return JsonResponse({'success': False, 'message': 'Método inválido'}, status=400)


def obter_itens_disponiveis():
    """
    Função auxiliar que retorna o dicionário de itens disponíveis por categoria.
    """
    itens_disponiveis = Subcategoria.objects.filter(situacao='disponivel')

    itens_por_categoria = {}
    for item in itens_disponiveis:
        categoria = item.categoria
        if categoria not in itens_por_categoria:
            itens_por_categoria[categoria] = []
        itens_por_categoria[categoria].append(item)

    # Imprime detalhes dos itens no terminal
    for categoria, itens in itens_por_categoria.items():
        print(f"Categoria: {categoria}")
        for item in itens:
            print(f"Nome: {item.descricao_completa}")
            print(f"Marca: {item.marca}")
            print(f"Modelo: {item.modelo}")
            print(f"Calibre: {item.cal}")
            print(f"Nº Arma: {item.num_arma}")
            print(f"Nº PMMA: {item.num_pmma}")
            print(f"Localização: {item.localizacao}")
            print(f"Estado de Conservação: {item.estado_conservacao}")
            print(f"Observação: {item.observacao}")
            print("----------------------------------------")

    return itens_por_categoria


def itens_disponiveis(request):
    # Filtra todos os itens que estão marcados como 'disponível'
    itens_disponiveis = Subcategoria.objects.filter(situacao='disponivel')

    # Agrupa os itens por categoria
    itens_por_categoria = {}
    for item in itens_disponiveis:
        # Assumindo que 'categoria' é um campo relacionado ao modelo Categoria
        categoria = item.categoria
        if categoria not in itens_por_categoria:
            itens_por_categoria[categoria] = []
        itens_por_categoria[categoria].append(item)

    itens_por_categoria = obter_itens_disponiveis()  # Obtém os itens disponíveis
    return render(request, 'cautelaarmamento/templates/catalogo_de_equipamento/itens_disponiveis.html', {
        'itens_por_categoria': itens_por_categoria
    })


def listar_inventario_equipamentos(request):
    # Filtra todos os itens disponíveis, ordenando-os por categoria
    itens_disponiveis = Subcategoria.objects.all().order_by('categoria')

    # Debug: Imprime os detalhes dos itens no terminal, organizados por categoria
    # for item in itens_disponiveis:
    #     print(f"Categoria: {item.categoria}")
    #     print(f"Nome: {item.descricao_completa}")
    #     print(f"Marca: {item.marca}")
    #     print(f"Modelo: {item.modelo}")
    #     print(f"Calibre: {item.cal}")
    #     print(f"Nº Arma: {item.num_arma}")
    #     print(f"Nº PMMA: {item.num_pmma}")
    #     print(f"Localização: {item.localizacao}")
    #     print(f"Estado de Conservação: {item.estado_conservacao}")
    #     print(f"Observação: {item.observacao}")
    #     print("----------------------------------------")

    # Renderiza o template com a lista de itens disponíveis organizados por categoria
    return render(request, 'cautelaarmamento/templates/catalogo_de_equipamento/inventario_equipamentos.html', {
        'itens_disponiveis': itens_disponiveis,
    })


def substituir_marcadores(doc, substituicoes):
    # Percorrer todos os parágrafos e substituir os marcadores
    for p in doc.paragraphs:
        for marcador, valor in substituicoes.items():
            if valor is None:
                valor = ""  # Substitui None por uma string vazia ou algum valor padrão
            if marcador in p.text:
                # Substituir o marcador pelo valor
                inline = p.runs
                for i in inline:
                    if marcador in i.text:
                        i.text = i.text.replace(marcador, valor)

def registrar_passagem(request):
    if request.method == 'POST':
        registro_cautela_id = request.POST.get('registro_cautela_id')
        data_inicio_str = request.POST.get('dataInicio')
        nome_substituto = request.POST.get('nomeSubstituto')
        observacoes = request.POST.get('observacoes')
        hora_atual = datetime.now().strftime('%H:%M')
        nome_usuario = request.user.first_name  # Isso retorna o primeiro nome do usuário autenticado

        try:
            data_inicio = timezone.make_aware(datetime.strptime(data_inicio_str, '%Y-%m-%d'))
        except (ValueError, TypeError):
            return render(request, 'passagem_de_servico/registrar_passagem.html', {
                'error': 'Data de início inválida',
                'registro_cautela': registro_cautela_id,
                'usuarios': User.objects.all(),
                'dataInicio': data_inicio_str,
                'nomeSubstituto': nome_substituto,
                'hora_atual': hora_atual,
                'observacoes': observacoes,
                'registros': PassagemDeServico.objects.all()
            })

        data_fim = timezone.now()
        
        cautelas_queryset = CautelaDeArmamento.objects.filter(
            hora_cautela__range=(data_inicio, data_fim),
            armeiro=request.user
        )
        
        cautelas_de_municoes_queryset = CautelaDeMunicoes.objects.filter(
            data_descautelamento__range=(data_inicio, data_fim),  # Corrigido para 'data_descautelamento'
            armeiro=request.user
        )

        descautelas_queryset = RegistroDescautelamento.objects.filter(
            data_descautelamento__range=(data_inicio, data_fim),
            armeiro=request.user
        )
        
        descautelas_ca_queryset = DescautelasCa.objects.filter(
            data_descautelamento__range=(data_inicio, data_fim),
            armeiro=request.user
        )

        # Obter itens disponíveis
        itens_por_categoria = obter_itens_disponiveis()
        
        # material_disponivel_armamento_queryset = Subcategoria.objects.filter(
        #     disponivel=True,
        #     armeiro=request.user
        # )
        
        # material_disponivel_municoes_queryset = CategoriaMunicao.objects.filter(
        #     data_descautelamento__range=(data_inicio, data_fim),
        #     armeiro=request.user
        # )
        
        # Caminho absoluto da imagem do logotipo
        logotipo_path = os.path.join(settings.BASE_DIR, 'cautelaarmamento', 'static', 'cautelaarmamento', 'images', 'logo.png')

        # Verifique se o logotipo existe
        if not os.path.exists(logotipo_path):
            raise FileNotFoundError(f"Erro: O logotipo não foi encontrado no caminho: {logotipo_path}")

        
        # Gerar conteúdo HTML para o relatório
        html_content = render_to_string('passagem_de_servico/relatorio.html', {
            'usuario': request.user,
            'hora_atual': hora_atual,
            'nome_substituto': nome_substituto,
            'data_inicio': data_inicio,
            'data_fim': data_fim,
            'cautelas': cautelas_queryset,
            'cautela_municoes':cautelas_de_municoes_queryset,
            'descautelas': descautelas_queryset,
            'descautela_ca':descautelas_ca_queryset,
            'material_disponivel': itens_por_categoria,
            'observacoes': observacoes,
            
        })

        # Salvar o conteúdo HTML em um arquivo na pasta 'relatorios'
        relatorios_path = os.path.join(settings.BASE_DIR, 'relatorios')
        os.makedirs(relatorios_path, exist_ok=True)  # Cria a pasta caso não exista

        filename = f'relatorio_{data_fim.strftime("%Y%m%d_%H%M%S")}.html'
        file_path = os.path.join(relatorios_path, filename)

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
#############################################################

        # Garantir que os dados recebidos não sejam None
        usuario = request.user.username  # Obtendo o nome do usuário logado
        nome_substituto = request.POST.get('nomeSubstituto', 'Não especificado')  # Valor padrão
        data_inicio_str = request.POST.get('dataInicio', 'Não especificado')  # Valor padrão
        hora_atual = datetime.now().strftime('%H:%M')
        data_fim = timezone.now()

        # Caminho para o documento modelo
        docx_path = 'relatorios/RELATORIO.docx'  # Caminho correto para o documento modelo
        doc = Document(docx_path)

        # Dados para substituir os marcadores
        substituicoes = {
            '<<ARMEIRO>>': usuario,  # Agora irá pegar o nome do usuário logado
            '<<NOME>>': nome_substituto,
            '<<DATA_INICIO>>': data_inicio_str,
            '<<HORA_ATUAL>>': hora_atual,
            '<<DATA_FIM>>': data_fim.strftime('%d/%m/%Y'),
        }

        # Criar a tabela para "Cautela de Cautelas"
        # Adicionar título para a tabela de cautelas
        doc.add_paragraph('Tabela de Cautelas', style='Título 12')  # Título acima da tabela, 'Heading 2'

        table = doc.add_table(rows=1, cols=4)

        # Cabeçalhos da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Policial'
        hdr_cells[2].text = 'Subcategoria'
        hdr_cells[3].text = 'Data da Cautela'

        # Ajustando a largura de cada coluna
        # A soma total das larguras não pode ultrapassar 6.5 polegadas (a largura padrão de uma página A4)
        table_width = 8.0  # largura total da tabela (em polegadas)

        # Proporção para ajustar as larguras das colunas
        col1_width = 0.5  # Largura da coluna 'ID' (em polegadas)
        col2_width = 2.0  # Largura da coluna 'Policial' (em polegadas)
        col3_width = 3.5  # Largura da coluna 'Subcategoria' (em polegadas)
        col4_width = 1.5  # Largura da coluna 'Data da Cautela' (em polegadas)

        # Definir as larguras das colunas de forma proporcional
        table.columns[0].width = Inches(col1_width)
        table.columns[1].width = Inches(col2_width)
        table.columns[2].width = Inches(col3_width)
        table.columns[3].width = Inches(col4_width)

        # Verificar se a soma das larguras não ultrapassa 6.5 polegadas
        total_width = col1_width + col2_width + col3_width + col4_width
        if total_width > table_width:
            print(f"A soma das larguras das colunas ({total_width} polegadas) ultrapassa a largura da página ({table_width} polegadas). Ajuste as larguras das colunas.")
        else:
            print(f"As larguras das colunas estão dentro do limite da página. Soma total: {total_width} polegadas.")

        # Preencher a tabela com os dados de 'cautelas_queryset'
        for cautela in cautelas_queryset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cautela.id)
            row_cells[1].text = str(cautela.policial)
            row_cells[2].text = str(cautela.subcategoria)
            row_cells[3].text = cautela.hora_cautela.strftime("%d/%m/%Y %H:%M")

            # Ajustando a largura de cada célula para cada linha adicionada
            row_cells[0]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="500" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[1]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="2500" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[2]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="5000" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[3]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="2000" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

            # Definir a borda de cada célula para ter uma linha visível
            for cell in row_cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'  <w:top w:val="single" w:sz="4" />'
                    r'  <w:left w:val="single" w:sz="4" />'
                    r'  <w:bottom w:val="single" w:sz="4" />'
                    r'  <w:right w:val="single" w:sz="4" />'
                    r'</w:tcBorders>'
                )
                tcPr.append(tcBorders)

        # Adicionar bordas a todas as células da tabela
        for row in table.rows:
            for cell in row.cells:
                # Definir a borda de cada célula para ter uma linha visível
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'  <w:top w:val="single" w:sz="4" />'
                    r'  <w:left w:val="single" w:sz="4" />'
                    r'  <w:bottom w:val="single" w:sz="4" />'
                    r'  <w:right w:val="single" w:sz="4" />'
                    r'</w:tcBorders>'
                )
                tcPr.append(tcBorders)


        # Adicionar a segunda tabela "Cautela de Munições"
        # Título para a tabela de Cautela de Munições
        doc.add_paragraph('Cautela de Munições', style='Título 12')  # Título acima da tabela de munições

        # Criar a tabela
        table = doc.add_table(rows=1, cols=4)

        # Cabeçalhos da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Policial'
        hdr_cells[2].text = 'Subcategoria'
        hdr_cells[3].text = 'Quantidade'

        # Ajustando a largura de cada coluna
        # A soma total das larguras não pode ultrapassar 6.5 polegadas (a largura padrão de uma página A4)
        table_width = 8.0  # largura total da tabela (em polegadas)

        # Proporção para ajustar as larguras das colunas
        col1_width = 0.5  # Largura da coluna 'ID' (em polegadas)
        col2_width = 2.0  # Largura da coluna 'Policial' (em polegadas)
        col3_width = 3.5  # Largura da coluna 'Subcategoria' (em polegadas)
        col4_width = 1.5  # Largura da coluna 'Quantidade' (em polegadas)

        # Definindo as larguras das colunas
        table.columns[0].width = Inches(col1_width)
        table.columns[1].width = Inches(col2_width)
        table.columns[2].width = Inches(col3_width)
        table.columns[3].width = Inches(col4_width)

        # Verificar se a soma das larguras não ultrapassa 6.5 polegadas
        total_width = col1_width + col2_width + col3_width + col4_width
        if total_width > table_width:
            print(f"A soma das larguras das colunas ({total_width} polegadas) ultrapassa a largura da página ({table_width} polegadas). Ajuste as larguras das colunas.")
        else:
            print(f"As larguras das colunas estão dentro do limite da página. Soma total: {total_width} polegadas.")

        # Preencher a tabela com os dados de 'cautela_municoes'
        for cautela in cautelas_de_municoes_queryset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cautela.id)
            row_cells[1].text = str(cautela.policial)
            row_cells[2].text = str(cautela.subcategoria)
            row_cells[3].text = str(cautela.quantidade)

            # Ajustar a largura de cada célula para cada linha adicionada (se necessário)
            row_cells[0]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="200" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[1]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="3000" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[2]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="8000" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[3]._element.get_or_add_tcPr().append(parse_xml('<w:tcW w:w="2000" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Definir a borda de cada célula para ter uma linha visível
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'  <w:top w:val="single" w:sz="4" />'
                    r'  <w:left w:val="single" w:sz="4" />'
                    r'  <w:bottom w:val="single" w:sz="4" />'
                    r'  <w:right w:val="single" w:sz="4" />'
                    r'</w:tcBorders>'
                )
                tcPr.append(tcBorders)


        # Adicionar a segunda tabela "Descautelas S/A"
        doc.add_paragraph('Descautelas S/A', style='Título 12')  # Título acima da tabela de descautelas

        # Criar a tabela
        table = doc.add_table(rows=1, cols=5)  # Adicionando uma coluna a mais para 'hora_descautelamento'

        # Cabeçalhos da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Policial'
        hdr_cells[2].text = 'Subcategoria Armamento'  # Agora estamos usando o campo subcategoria_armamento
        hdr_cells[3].text = 'Quantidade'
        hdr_cells[4].text = 'Data/Hora'  # Nova coluna para hora de descautelamento

        # Definir as larguras das colunas
        col1_width = Inches(0.5)  # Largura da coluna 'ID'
        col2_width = Inches(1.8)  # Largura da coluna 'Policial'
        col3_width = Inches(3.0)  # Largura da coluna 'Subcategoria Armamento' (ajustada)
        col4_width = Inches(1.2)  # Largura da coluna 'Quantidade' (ajustada)
        col5_width = Inches(1.0)  # Largura da coluna 'Data/Hora'

        # Calcular a largura total e garantir que não ultrapasse os limites
        total_width = col1_width + col2_width + col3_width + col4_width + col5_width
        if total_width > Inches(8.0):  # Ajuste para não ultrapassar a largura da página
            col3_width = Inches(2.0)  # Ajuste de largura da coluna 'Subcategoria Armamento'
            col4_width = Inches(1.0)  # Ajuste de largura da coluna 'Quantidade'

        # Aplicar as larguras para as colunas
        table.columns[0].width = col1_width
        table.columns[1].width = col2_width
        table.columns[2].width = col3_width
        table.columns[3].width = col4_width
        table.columns[4].width = col5_width

        # Preencher a tabela com os dados de 'descautelas'
        for cautela in descautelas_queryset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cautela.id)
            row_cells[1].text = str(cautela.policial)
            row_cells[2].text = str(cautela.subcategoria_armamento)  # Usando o campo correto
            row_cells[3].text = str(cautela.quantidade_municao)  # Usando quantidade de munição
            row_cells[4].text = str(cautela.data_descautelamento)  # Usando a hora de descautelamento

            # Ajustar a largura de cada célula para cada linha adicionada
            row_cells[0]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col1_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[1]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col2_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[2]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col3_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[3]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col4_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[4]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col5_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Definir a borda de cada célula para ter uma linha visível
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'  <w:top w:val="single" w:sz="4" />'
                    r'  <w:left w:val="single" w:sz="4" />'
                    r'  <w:bottom w:val="single" w:sz="4" />'
                    r'  <w:right w:val="single" w:sz="4" />'
                    r'</w:tcBorders>'
                )
                tcPr.append(tcBorders)




        # Adicionar a segunda tabela "Cautela de Munições"
        # Título para a tabela de Cautela de Munições
        doc.add_paragraph('Descautela C/A', style='Título 12')  # Título acima da tabela de munições

        # Criar a tabela
        table = doc.add_table(rows=1, cols=7)  # Adicionando uma coluna a mais para 'hora_descautelamento'

        # Cabeçalhos da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Policial'
        hdr_cells[2].text = 'Armamento'  # Agora estamos usando o campo subcategoria_armamento
        hdr_cells[3].text = 'Munição'
        hdr_cells[4].text = 'Quant. Munição'
        hdr_cells[5].text = 'Situação'
        hdr_cells[6].text = 'Observação'

        # Definir as larguras das colunas considerando as margens da página
        col1_width = Inches(0.5)  # Largura da coluna 'ID'
        col2_width = Inches(2.0)  # Largura da coluna 'Policial'
        col3_width = Inches(3.0)  # Largura da coluna 'Armamento'
        col4_width = Inches(1.2)  # Largura da coluna 'Munição'
        col5_width = Inches(1.0)  # Largura da coluna 'Quant. Munição'
        col6_width = Inches(1.2)  # Largura da coluna 'Situação'
        col7_width = Inches(1.5)  # Largura da coluna 'Observação'

        # Calcular a largura total e garantir que não ultrapasse os limites da página
        total_width = col1_width + col2_width + col3_width + col4_width + col5_width + col6_width + col7_width
        max_width = Inches(6.5)  # Largura útil da página (ajustar conforme as margens do documento)

        # Ajustar as larguras se necessário para garantir que a soma total não ultrapasse a largura útil
        if total_width > max_width:  # Ajuste para não ultrapassar a largura da página
            factor = max_width / total_width  # Fator de escala para ajustar a largura das colunas
            col1_width *= factor
            col2_width *= factor
            col3_width *= factor
            col4_width *= factor
            col5_width *= factor
            col6_width *= factor
            col7_width *= factor

            # Aplicar as larguras para as colunas, convertendo os valores para int
            table.columns[0].width = int(col1_width * 1440)  # Multiplica por 1440 para converter para twips (unidade esperada)
            table.columns[1].width = int(col2_width * 1440)
            table.columns[2].width = int(col3_width * 1440)
            table.columns[3].width = int(col4_width * 1440)
            table.columns[4].width = int(col5_width * 1440)
            table.columns[5].width = int(col6_width * 1440)
            table.columns[6].width = int(col7_width * 1440)

        # Preencher a tabela com os dados de 'descautelas'
        for cautela in descautelas_ca_queryset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cautela.id)
            row_cells[1].text = str(cautela.policial)
            row_cells[2].text = str(cautela.subcategoria_armamento)  # Usando o campo correto
            row_cells[3].text = str(cautela.subcategoria_municao)  # Usando quantidade de munição
            row_cells[4].text = str(cautela.quantidade_municao)  # Usando a quantidade de munição
            row_cells[5].text = str(cautela.situacao_armamento) 
            row_cells[6].text = str(cautela.observacao)

            # Ajustar a largura de cada célula para cada linha adicionada
            row_cells[0]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col1_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[1]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col2_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[2]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col3_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[3]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col4_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[4]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col5_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[5]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col6_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            row_cells[6]._element.get_or_add_tcPr().append(parse_xml(f'<w:tcW w:w="{int(col7_width * 1440)}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        # Definir a borda de cada célula para ter uma linha visível
        for row in table.rows:
            for cell in row.cells:
                # Definir a borda de cada célula
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'  <w:top w:val="single" w:sz="4" />'
                    r'  <w:left w:val="single" w:sz="4" />'
                    r'  <w:bottom w:val="single" w:sz="4" />'
                    r'  <w:right w:val="single" w:sz="4" />'
                    r'</w:tcBorders>'
                )
                tcPr.append(tcBorders)

        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Adicionar o título da tabela
        titulo = doc.add_paragraph('MATERIAL DISPONÍVEL NA RESERVA', style='Título 12')
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centralizar o título

        # Criar a tabela com a quantidade adequada de colunas (apenas 3 colunas para dados)
        table = doc.add_table(rows=1, cols=3)

        # Definir o estilo da tabela
        table.style = 'Table Grid'

        # Cabeçalhos da tabela (somente 3 colunas)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Descrição'
        hdr_cells[1].text = 'Estado de Conservação'
        hdr_cells[2].text = 'Observação'

        # Preencher a tabela com os dados de 'material_disponivel'
        if itens_por_categoria:  # Verificar se há categorias
            for categoria, itens in itens_por_categoria.items():
                # Adicionar uma linha para a categoria (sem mesclar células)
                categoria_row = table.add_row().cells
                categoria_row[0].text = str(categoria).upper()  # Exibir a categoria em maiúsculas
                categoria_row[0].paragraphs[0].runs[0].bold = True  # Destacar a categoria em negrito
                categoria_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                categoria_row[0].paragraphs[0].paragraph_format.space_before = Pt(6)
                categoria_row[0].paragraphs[0].paragraph_format.space_after = Pt(6)

                # Adicionar os itens da categoria
                for item in itens:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.descricao_completa)
                    row_cells[1].text = str(item.estado_conservacao)
                    row_cells[2].text = str(item.observacao)
        else:
            # Caso não haja itens disponíveis, adicionar uma mensagem de aviso
            no_item_paragraph = doc.add_paragraph('Nenhum item disponível na reserva.', style='Normal')
            no_item_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph('Fiz ao meu substituto legal o(a) <<NOME>> ,a quem repassei todas as ordens em vigor bem como todo o material da carga da 4ªCIA/11ºBPM conforme os dados acima', style='Normal')

        # Substituindo os marcadores no documento
        substituir_marcadores(doc, substituicoes)

        # Salvar o arquivo .docx modificado
        docx_path = os.path.join(settings.BASE_DIR, 'relatorios', 'RELATORIO1.docx')
        doc.save(docx_path)

        # Caminho de saída do PDF
        output_path = os.path.join(settings.BASE_DIR, 'relatorios', 'RELATORIO.pdf')

        # Convertendo de .docx para .pdf usando LibreOffice no modo de linha de comando
        try:
            # Comando para converter .docx para .pdf usando LibreOffice
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', os.path.dirname(output_path)], check=True)
            print(f"PDF gerado com sucesso: {output_path}")
        except subprocess.CalledProcessError as e:
            print(f"Erro ao gerar PDF: {e}")


######################################################

        # Retornar o arquivo HTML como resposta
        with open(file_path, 'r', encoding='utf-8') as f:
            response = HttpResponse(f.read(), content_type='text/html')

        # Opcional: Retornar o arquivo .docx para download
        with open(docx_path, 'rb') as f:
            response_docx = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response_docx['Content-Disposition'] = f'attachment; filename="relatorio_passagem_{data_fim.strftime("%Y%m%d_%H%M%S")}.docx"'

        # Retornar ambos os arquivos se necessário
        return response  # Ou retornar 'response_docx' para o arquivo .docx
    else:
        usuarios = User.objects.all()
        registros = PassagemDeServico.objects.all().order_by('data_inicio')

        paginator = Paginator(registros, 7)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        return render(request, 'passagem_de_servico/registrar_passagem.html', {
            'registro_cautela': None,
            'usuarios': usuarios,
            'dataInicio': '',
            'nomeSubstituto': '',
            'observacoes': '',
            'page_obj': page_obj
        })


def gerar_relatorio(request):
    if request.method == 'POST':
        # Coleta os dados necessários
        passagem = PassagemDeServico.objects.filter(usuario=request.user).latest('data_inicio')
        data_fim = passagem.data_fim
        data_inicio = passagem.data_inicio
        hora_atual = timezone.now().strftime('%H:%M')
        
        # Consultas para coletar os dados do relatório
        cautelas = CautelaDeArmamento.objects.filter(hora_cautela__range=(data_inicio, data_fim), armeiro=request.user)
        descautelas = RegistroDescautelamento.objects.filter(data_descautelamento__range=(data_inicio, data_fim), armeiro=request.user)
        cautela_de_municoes = CautelaDeMunicoes.objects.filter(data_descautelamento__range=(data_inicio, data_fim), armeiro=request.user)
        descautelas_ca = DescautelasCa.objects.filter(data_descautelamento__range=(data_inicio, data_fim), armeiro=request.user)
        
        # Obter itens disponíveis
        itens_por_categoria = obter_itens_disponiveis()

        # Caminho do arquivo HTML na pasta 'relatorios'
        html_file_path = os.path.join(settings.BASE_DIR, 'relatorios', 'relatorio_template.html')
        
        # Lê o conteúdo do arquivo HTML
        with open(html_file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        # Substituir os placeholders no HTML pelos dados dinâmicos
        html_content = html_content.replace("{{ usuario }}", request.user.username)
        html_content = html_content.replace("{{ hora_atual }}", hora_atual)
        html_content = html_content.replace("{{ nome_substituto }}", passagem.nome_substituto)
        html_content = html_content.replace("{{ data_inicio }}", data_inicio.strftime('%d/%m/%Y'))
        html_content = html_content.replace("{{ data_fim }}", data_fim.strftime('%d/%m/%Y'))

        # Substituir as listas de cautelas e descautelas, como exemplo
        cautelas_html = ''.join([f"<tr><td>{c.id}</td><td>{c.policial}</td><td>{c.subcategoria}</td><td>{c.hora_cautela.strftime('%d/%m/%Y %H:%M')}</td></tr>" for c in cautelas])
        html_content = html_content.replace("{{ cautelas }}", cautelas_html)

        # Gerar o caminho para salvar o PDF
        pdf_file_path = os.path.join(settings.BASE_DIR, 'relatorios', f'relatorio_{data_fim.strftime("%Y%m%d_%H%M%S")}.pdf')
        
        # Usar o WeasyPrint para converter HTML para PDF
        html = HTML(string=html_content)

        # Configurar margens e formato da página A4
        html.write_pdf(pdf_file_path, stylesheets=[os.path.join(settings.BASE_DIR, 'static/css/styles.css')], 
                       presentational_hints=True, 
                       zoom=1, 
                       options={"margin-top": "1cm", "margin-right": "1cm", "margin-bottom": "1cm", "margin-left": "1cm"})
        
        # Retornar o PDF gerado
        with open(pdf_file_path, 'rb') as f:
            pdf_content = f.read()
        
        # Retornar o PDF como resposta
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = f'inline; filename="relatorio_{data_fim.strftime("%Y%m%d_%H%M%S")}.pdf"'
        
        return response

    else:
        # Retornar um erro ou redirecionar em caso de acesso GET não esperado
        return HttpResponse(status=405)  # Método não permitido


def listar_cautelas(request):
    # Obtém todas as instâncias de CautelaDeArmamento do banco de dados
    cautelas = CautelaDeArmamento.objects.all()


# Renderiza os dados no template
    return render(request, 'passagem_de_servico/listar_amas_cauteladas.html', {'cautelas': cautelas})


def listar_descautela(request):
    # Obtém todas as instâncias de CautelaDeArmamento do banco de dados
    descautelas = RegistroDescautelamento.objects.all()

    # Imprimir o dicionário no terminal
    print({'descautelas': descautelas})

    # Renderiza os dados no template
    return render(request, 'passagem_de_servico/listar_armas_descauteladas.html', {'descautelas': descautelas})


def materiais_alterados(request):
    # Filtra os materiais da subcategoria cuja situação NÃO seja "disponível" ou "cautelado"
    materiais_alterados = Subcategoria.objects.exclude(situacao='disponivel').exclude(situacao='cautelada')
    
    # Retorna uma resposta renderizando uma template com os dados
    return render(request, 'manutencao_e_suporte/listar_materiais_alterados.html', {'materiais_alterados': materiais_alterados})

def restaurar_status(request, material_id):
    # Recupera o material com base no ID
    material = get_object_or_404(Subcategoria, id=material_id)
    
    # Atualiza a situação do material para 'disponivel'
    material.situacao = 'disponivel'
    material.save()

    # Redireciona para a página de materiais alterados
    return redirect('listar_materiais_alterados')


def verificador_de_autorizacao(request):
    # Obtém todos os policiais
    policiais = Policial.objects.all()

    # Renderiza o template passando a lista de todos os policiais
    return render(request, 'passagem_de_servico/verificar_autorizacao.html', {'policiais': policiais})
